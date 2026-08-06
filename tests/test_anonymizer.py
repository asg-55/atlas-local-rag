from __future__ import annotations

import io
import json
import unittest
import zipfile
from email import policy
from email.message import EmailMessage
from email.parser import BytesParser

from docx import Document
from openpyxl import Workbook, load_workbook

from rag_assistant.anonymizer import (
    Finding,
    anonymize_document,
    find_sensitive_data,
    restore_document,
    xlsx_technical_columns,
)


PASSWORD = "correct horse battery staple"


def _docx_bytes(text: str) -> bytes:
    document = Document()
    paragraph = document.add_paragraph()
    midpoint = len(text) // 2
    paragraph.add_run(text[:midpoint])
    paragraph.add_run(text[midpoint:])
    target = io.BytesIO()
    document.save(target)
    return target.getvalue()


def _docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _xlsx_bytes(text: str) -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet["A1"] = text
    target = io.BytesIO()
    workbook.save(target)
    return target.getvalue()


def _pptx_bytes(text: str) -> bytes:
    slide = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">'
        f"<p:cSld><a:p><a:r><a:t>{text}</a:t></a:r></a:p></p:cSld></p:sld>"
    )
    target = io.BytesIO()
    with zipfile.ZipFile(target, "w") as package:
        package.writestr("ppt/slides/slide1.xml", slide)
    return target.getvalue()


def _pptx_text(content: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(content)) as package:
        return package.read("ppt/slides/slide1.xml").decode("utf-8")


class AnonymizerTests(unittest.TestCase):
    def test_docx_round_trip_across_runs_and_encrypted_key(self):
        original = "ФИО: Иванов Иван Иванович, email ivan@example.ru"
        content = _docx_bytes(original)
        findings = find_sensitive_data(content, "staff.docx", {"PERSON", "EMAIL"})

        self.assertEqual({item.category for item in findings}, {"PERSON", "EMAIL"})
        result = anonymize_document(content, "staff.docx", findings, PASSWORD)
        anonymous_text = _docx_text(result.content)
        self.assertIn("ATLAS-PERSON-0001", anonymous_text)
        self.assertIn("atlas-email-0001@anonymous.invalid", anonymous_text)
        self.assertNotIn("Иванов", anonymous_text)
        self.assertNotIn(b"ivan@example.ru", result.key_content)
        self.assertEqual(json.loads(result.key_content)["kdf"], "argon2id")

        restored = restore_document(result.content, result.filename, result.key_content, PASSWORD)
        self.assertEqual(_docx_text(restored.content), original)

    def test_xlsx_round_trip_preserves_workbook(self):
        original = "Телефон +7 (913) 123-45-67"
        content = _xlsx_bytes(original)
        findings = find_sensitive_data(content, "contacts.xlsx", {"PHONE"})
        result = anonymize_document(content, "contacts.xlsx", findings, PASSWORD)
        workbook = load_workbook(io.BytesIO(result.content))
        self.assertIn("ATLAS-PHONE-0001", workbook.active["A1"].value)

        restored = restore_document(result.content, result.filename, result.key_content, PASSWORD)
        workbook = load_workbook(io.BytesIO(restored.content))
        self.assertEqual(workbook.active["A1"].value, original)

    def test_xlsx_discovers_name_columns_without_touching_values_or_formulas(self):
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Выгрузка"
        worksheet.append(["Узел", "Тег", "Наименование параметра", "Значение", "Единица"])
        worksheet.append(
            [
                "Установка полимеризации",
                "FIC-1025",
                "Расход пропилена — Установка полимеризации",
                123.45,
                "кг/ч",
            ]
        )
        worksheet.append(
            ["Насос катализатора", "P-201A", "Давление насоса P-201A", "=D2*2", "МПа"]
        )
        source = io.BytesIO()
        workbook.save(source)
        content = source.getvalue()

        columns = xlsx_technical_columns(content)
        suggested = [column for column in columns if column.suggested]
        self.assertEqual({column.header for column in suggested}, {"Узел", "Тег"})
        findings = find_sensitive_data(
            content,
            "export.xlsx",
            {"TECHNICAL_TAG", "TECHNICAL_NAME"},
            technical_columns=[column.key for column in suggested],
        )
        values = {finding.value for finding in findings}
        self.assertTrue(
            {"Установка полимеризации", "Насос катализатора", "FIC-1025", "P-201A"}.issubset(
                values
            )
        )
        self.assertNotIn("Расход пропилена — Установка полимеризации", values)
        self.assertNotIn("Давление насоса P-201A", values)

        result = anonymize_document(content, "export.xlsx", findings, PASSWORD)
        anonymous = load_workbook(io.BytesIO(result.content), data_only=False)
        self.assertEqual(anonymous["Выгрузка"]["D2"].value, 123.45)
        self.assertEqual(anonymous["Выгрузка"]["D3"].value, "=D2*2")
        self.assertEqual(anonymous["Выгрузка"]["E2"].value, "кг/ч")
        self.assertIn("ATLAS-TECHNICAL", anonymous["Выгрузка"]["A2"].value)
        self.assertRegex(
            anonymous["Выгрузка"]["C2"].value,
            r"^Расход пропилена — ATLAS-TECHNICAL_NAME-\d+$",
        )
        self.assertRegex(
            anonymous["Выгрузка"]["C3"].value,
            r"^Давление насоса ATLAS-TECHNICAL_NAME-\d+$",
        )

        restored = restore_document(result.content, result.filename, result.key_content, PASSWORD)
        restored_book = load_workbook(io.BytesIO(restored.content), data_only=False)
        self.assertEqual(restored_book["Выгрузка"]["A2"].value, "Установка полимеризации")
        self.assertEqual(restored_book["Выгрузка"]["B2"].value, "FIC-1025")
        self.assertEqual(
            restored_book["Выгрузка"]["C2"].value,
            "Расход пропилена — Установка полимеризации",
        )
        self.assertEqual(restored_book["Выгрузка"]["C3"].value, "Давление насоса P-201A")
        self.assertEqual(restored_book["Выгрузка"]["D3"].value, "=D2*2")

    def test_technical_tag_pattern_does_not_select_plain_numbers(self):
        content = _docx_bytes(
            "Аппараты D228, УПП-2, TIR_4_123, TR-3-4-9, PIRA-2-172/4, "
            "TRA_G202, B-TIR-05A, QIRA_101_5.PV и 18 TIR 75D. "
            "Давление 12.5, температура 80, ГОСТ 123."
        )
        findings = find_sensitive_data(
            content,
            "scheme.docx",
            {"TECHNICAL_TAG"},
            detect_technical_tags=True,
        )
        self.assertEqual(
            {finding.value for finding in findings},
            {
                "D228",
                "УПП-2",
                "TIR_4_123",
                "TR-3-4-9",
                "PIRA-2-172/4",
                "TRA_G202",
                "B-TIR-05A",
                "QIRA_101_5.PV",
                "18 TIR 75D",
            },
        )

    def test_pptx_and_eml_are_supported(self):
        presentation = _pptx_bytes("Связаться: analyst@example.ru")
        ppt_findings = find_sensitive_data(presentation, "brief.pptx", {"EMAIL"})
        ppt_result = anonymize_document(presentation, "brief.pptx", ppt_findings, PASSWORD)
        self.assertNotIn("analyst@example.ru", _pptx_text(ppt_result.content))
        ppt_restored = restore_document(
            ppt_result.content, ppt_result.filename, ppt_result.key_content, PASSWORD
        )
        self.assertIn("analyst@example.ru", _pptx_text(ppt_restored.content))

        message = EmailMessage()
        message["From"] = "Иван Иванов <sender@example.ru>"
        message["To"] = "receiver@example.ru"
        message["Subject"] = "Контакт +7 913 111-22-33"
        message.set_content("Ответить receiver@example.ru")
        email_content = message.as_bytes(policy=policy.default)
        eml_findings = find_sensitive_data(email_content, "message.eml", {"EMAIL", "PHONE"})
        eml_result = anonymize_document(email_content, "message.eml", eml_findings, PASSWORD)
        parsed = BytesParser(policy=policy.default).parsebytes(eml_result.content)
        self.assertNotIn("receiver@example.ru", parsed.as_string())
        self.assertIn("@anonymous.invalid", str(parsed["From"]))
        eml_restored = restore_document(
            eml_result.content, eml_result.filename, eml_result.key_content, PASSWORD
        )
        restored_message = BytesParser(policy=policy.default).parsebytes(eml_restored.content)
        self.assertEqual(str(restored_message["From"]), "Иван Иванов <sender@example.ru>")
        self.assertEqual(str(restored_message["To"]), "receiver@example.ru")
        self.assertIn("receiver@example.ru", restored_message.get_content())

    def test_text_formats_round_trip_preserves_encoding_and_structure(self):
        cases = [
            (
                "signals.csv",
                "Тег;Почта\r\nTIR_4_123;operator@example.ru\r\n".encode("cp1251"),
                "cp1251",
            ),
            (
                "notes.txt",
                "Контур TR-3-4-9, ответственный operator@example.ru.\r\n".encode("utf-8"),
                "utf-8",
            ),
            (
                "payload.json",
                b"\xef\xbb\xbf" + '{\n  "tag": "QIRA_101_5.PV",\n  "email": "operator@example.ru"\n}\n'.encode("utf-8"),
                "utf-8-sig",
            ),
            (
                "readme.md",
                "# Сигнал\n\nПрибор **18 TIR 75D**, контакт operator@example.ru.\n".encode("utf-16"),
                "utf-16",
            ),
        ]
        for filename, content, encoding in cases:
            with self.subTest(filename=filename):
                findings = find_sensitive_data(
                    content,
                    filename,
                    {"EMAIL", "TECHNICAL_TAG"},
                )
                self.assertTrue(any(finding.category == "TECHNICAL_TAG" for finding in findings))
                self.assertTrue(any(finding.category == "EMAIL" for finding in findings))
                result = anonymize_document(content, filename, findings, PASSWORD)
                anonymous_text = result.content.decode(encoding)
                self.assertIn("ATLAS-TECHNICAL_TAG-", anonymous_text)
                self.assertIn("@anonymous.invalid", anonymous_text)
                if filename.endswith(".json"):
                    json.loads(anonymous_text)
                restored = restore_document(
                    result.content,
                    result.filename,
                    result.key_content,
                    PASSWORD,
                )
                self.assertEqual(restored.content, content)

    def test_office_key_restores_placeholders_in_processed_json(self):
        original = "TIR_4_123 operator@example.ru"
        content = _xlsx_bytes(original)
        findings = find_sensitive_data(content, "signals.xlsx", {"EMAIL", "TECHNICAL_TAG"})
        result = anonymize_document(content, "signals.xlsx", findings, PASSWORD)
        anonymous_book = load_workbook(io.BytesIO(result.content))
        anonymous_value = anonymous_book.active["A1"].value
        processed_json = json.dumps(
            {"answer": anonymous_value},
            ensure_ascii=False,
            indent=2,
        ).encode("utf-8")

        restored = restore_document(
            processed_json,
            "Replay_result.json",
            result.key_content,
            PASSWORD,
        )

        self.assertEqual(json.loads(restored.content)["answer"], original)
        self.assertEqual(restored.filename, "Replay_result_restored.json")
        with self.assertRaisesRegex(ValueError, "не совпадает с ключом"):
            restore_document(
                _docx_bytes(anonymous_value),
                "wrong.docx",
                result.key_content,
                PASSWORD,
            )

    def test_wrong_password_and_hostile_key_are_rejected(self):
        content = _docx_bytes("Почта secret@example.ru")
        result = anonymize_document(
            content,
            "secret.docx",
            [Finding("EMAIL", "secret@example.ru", 1, 0)],
            PASSWORD,
        )
        with self.assertRaisesRegex(ValueError, "Неверный пароль"):
            restore_document(result.content, result.filename, result.key_content, "another long password")

        hostile_key = json.loads(result.key_content)
        hostile_key["parameters"]["memory_cost"] = 2**31
        with self.assertRaisesRegex(ValueError, "повреждён|неизвестный формат"):
            restore_document(
                result.content,
                result.filename,
                json.dumps(hostile_key).encode("utf-8"),
                PASSWORD,
            )


if __name__ == "__main__":
    unittest.main()
