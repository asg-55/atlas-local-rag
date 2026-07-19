from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
from functools import lru_cache
from pathlib import Path

import cv2
import fitz
import numpy as np
import pandas as pd
from docx import Document as DocxDocument
from lxml import html as lxml_html
from PIL import Image

from .models import TextBlock
from .config import settings


@lru_cache(maxsize=1)
def _ocr_reader():
    import easyocr

    return easyocr.Reader(["ru", "en"], gpu=False, verbose=False)


@lru_cache(maxsize=1)
def _whisper_model():
    from faster_whisper import WhisperModel

    return WhisperModel("small", device="cpu", compute_type="int8")


def _box_metrics(box) -> tuple[float, float, float, float]:
    xs = [float(point[0]) for point in box]
    ys = [float(point[1]) for point in box]
    return min(xs), min(ys), max(xs), max(ys)


def _group_text_lines(items: list[dict]) -> list[str]:
    if not items:
        return []
    ordered = sorted(items, key=lambda item: (item["cy"], item["x1"]))
    rows: list[list[dict]] = []
    for item in ordered:
        if not rows:
            rows.append([item])
            continue
        current = rows[-1]
        row_y = sum(value["cy"] for value in current) / len(current)
        tolerance = max(10.0, item["height"] * 0.65)
        if abs(item["cy"] - row_y) <= tolerance:
            current.append(item)
        else:
            rows.append([item])
    rendered: list[str] = []
    for row in rows:
        row.sort(key=lambda item: item["x1"])
        parts: list[str] = []
        previous_x2 = None
        for item in row:
            if previous_x2 is not None:
                gap = item["x1"] - previous_x2
                separator = " | " if gap > max(35.0, item["height"] * 2.5) else " "
                parts.append(separator)
            parts.append(item["text"])
            previous_x2 = item["x2"]
        text = "".join(parts).strip()
        if text:
            rendered.append(text)
    return rendered


def _table_grid(image_array: np.ndarray) -> np.ndarray:
    gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
    binary = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 31, 12
    )
    height, width = gray.shape
    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(20, width // 35), 1))
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(20, height // 35)))
    horizontal = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel)
    vertical = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel)
    return cv2.bitwise_or(horizontal, vertical)


def _group_cell_rows(boxes: list[tuple[int, int, int, int]]) -> list[list[tuple[int, int, int, int]]]:
    rows: list[list[tuple[int, int, int, int]]] = []
    for box in sorted(boxes, key=lambda value: (value[1], value[0])):
        cy = (box[1] + box[3]) / 2
        matching = None
        for row in rows:
            row_cy = sum((item[1] + item[3]) / 2 for item in row) / len(row)
            median_height = np.median([item[3] - item[1] for item in row])
            if abs(cy - row_cy) <= max(8, median_height * 0.42):
                matching = row
                break
        if matching is None:
            rows.append([box])
        else:
            matching.append(box)
    valid_rows = []
    for row in rows:
        row.sort(key=lambda value: value[0])
        deduplicated = []
        for box in row:
            if any(
                abs(box[0] - old[0]) < 4 and abs(box[1] - old[1]) < 4
                and abs(box[2] - old[2]) < 4 and abs(box[3] - old[3]) < 4
                for old in deduplicated
            ):
                continue
            deduplicated.append(box)
        if len(deduplicated) >= 2:
            valid_rows.append(deduplicated)
    valid_rows.sort(key=lambda row: min(item[1] for item in row))
    return valid_rows


def _detect_table_regions(image_array: np.ndarray) -> list[list[list[tuple[int, int, int, int]]]]:
    grid = _table_grid(image_array)
    page_height, page_width = grid.shape
    external, _ = cv2.findContours(grid, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    regions = []
    for contour in external:
        x, y, w, h = cv2.boundingRect(contour)
        if w < page_width * 0.18 or h < 45 or w * h < page_width * page_height * 0.008:
            continue
        regions.append((x, y, x + w, y + h))
    regions.sort(key=lambda value: (value[1], value[0]))
    tables = []
    for region_x1, region_y1, region_x2, region_y2 in regions:
        crop = grid[region_y1:region_y2, region_x1:region_x2]
        region_height, region_width = crop.shape
        contours, _ = cv2.findContours(crop, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
        boxes: list[tuple[int, int, int, int]] = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            if w < 35 or h < 18 or w * h < 900:
                continue
            if w > region_width * 0.94 and h > region_height * 0.80:
                continue
            boxes.append((region_x1 + x, region_y1 + y, region_x1 + x + w, region_y1 + y + h))
        rows = _group_cell_rows(boxes)
        if len(rows) >= 2:
            tables.append(rows)
    return tables


def _detect_table_cells(image_array: np.ndarray) -> list[list[tuple[int, int, int, int]]]:
    """Compatibility helper returning rows from every detected table region."""
    return [row for table in _detect_table_regions(image_array) for row in table]


def _legacy_detect_table_cells(image_array: np.ndarray) -> list[list[tuple[int, int, int, int]]]:
    """Deprecated flat detector retained only to ease comparison during development."""
    grid = _table_grid(image_array)
    height, width = grid.shape
    contours, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    boxes: list[tuple[int, int, int, int]] = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if w < 35 or h < 18 or w > width * 0.96 or h > height * 0.85:
            continue
        if w * h < 900:
            continue
        boxes.append((x, y, x + w, y + h))
    unique: list[tuple[int, int, int, int]] = []
    for box in sorted(set(boxes), key=lambda value: (value[1], value[0], value[2] - value[0])):
        x1, y1, x2, y2 = box
        if any(abs(x1-a) < 4 and abs(y1-b) < 4 and abs(x2-c) < 4 and abs(y2-d) < 4 for a,b,c,d in unique):
            continue
        unique.append(box)
    valid_rows = _group_cell_rows(unique)
    return valid_rows if len(valid_rows) >= 2 else []


def _prepare_ocr_image(image: Image.Image) -> np.ndarray:
    image_array = np.array(image.convert("RGB"))
    height, width = image_array.shape[:2]
    if width < 1200:
        scale = min(2.5, 1200 / max(1, width))
        image_array = cv2.resize(image_array, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
    edges = cv2.Canny(gray, 50, 150, apertureSize=3)
    lines = cv2.HoughLinesP(
        edges,
        1,
        np.pi / 180,
        threshold=max(80, image_array.shape[1] // 8),
        minLineLength=max(100, image_array.shape[1] // 4),
        maxLineGap=20,
    )
    angles = []
    if lines is not None:
        for line in np.asarray(lines).reshape(-1, 4):
            x1, y1, x2, y2 = line
            angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
            if -12 <= angle <= 12:
                angles.append(angle)
    if angles:
        angle = float(np.median(angles))
        if 0.2 <= abs(angle) <= 12:
            h, w = image_array.shape[:2]
            matrix = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
            image_array = cv2.warpAffine(
                image_array,
                matrix,
                (w, h),
                flags=cv2.INTER_CUBIC,
                borderMode=cv2.BORDER_CONSTANT,
                borderValue=(255, 255, 255),
            )
    return image_array


def _ocr_image_blocks(image: Image.Image, location_prefix: str) -> list[TextBlock]:
    image_array = _prepare_ocr_image(image)
    raw = _ocr_reader().readtext(image_array, detail=1, paragraph=False)
    items: list[dict] = []
    for box, text, confidence in raw:
        text = " ".join(str(text).split())
        if not text:
            continue
        x1, y1, x2, y2 = _box_metrics(box)
        items.append(
            {
                "text": text,
                "confidence": float(confidence),
                "x1": x1,
                "y1": y1,
                "x2": x2,
                "y2": y2,
                "cx": (x1 + x2) / 2,
                "cy": (y1 + y2) / 2,
                "height": max(1.0, y2 - y1),
            }
        )
    tables = _detect_table_regions(image_array)
    table_boxes = [box for table in tables for row in table for box in row]
    blocks: list[TextBlock] = []
    for table_number, table_rows in enumerate(tables, start=1):
        rendered_rows = []
        confidences = []
        for row in table_rows:
            cells = []
            for x1, y1, x2, y2 in row:
                matches = [
                    item for item in items
                    if x1 - 3 <= item["cx"] <= x2 + 3 and y1 - 3 <= item["cy"] <= y2 + 3
                ]
                matches.sort(key=lambda item: (item["cy"], item["x1"]))
                cells.append(" ".join(item["text"] for item in matches).strip())
                confidences.extend(item["confidence"] for item in matches)
            if len(row) <= 3 and any(cells):
                for cell_index, ((x1, y1, x2, y2), cell_text) in enumerate(zip(row, cells)):
                    if cell_text:
                        continue
                    crop = image_array[max(0, y1 + 3) : max(0, y2 - 3), max(0, x1 + 3) : max(0, x2 - 3)]
                    if crop.size == 0:
                        continue
                    retry = _ocr_reader().readtext(
                        crop,
                        detail=0,
                        paragraph=True,
                        mag_ratio=2.0,
                        text_threshold=0.35,
                        low_text=0.15,
                    )
                    retried_text = " ".join(" ".join(str(value).split()) for value in retry).strip()
                    if retried_text:
                        cells[cell_index] = retried_text
            if any(cells):
                rendered_rows.append(" | ".join(cells))
        if rendered_rows:
            blocks.append(
                TextBlock(
                    text="\n".join(rendered_rows),
                    location=f"{location_prefix}, OCR-таблица {table_number}",
                    block_type="ocr_table",
                    metadata={"ocr": True, "table": True, "table_number": table_number, "confidence": round(float(np.mean(confidences)), 3) if confidences else 0.0},
                )
            )
    outside = []
    for item in items:
        inside_table = any(x1 <= item["cx"] <= x2 and y1 <= item["cy"] <= y2 for x1, y1, x2, y2 in table_boxes)
        if not inside_table:
            outside.append(item)
    lines = _group_text_lines(outside)
    if lines:
        confidences = [item["confidence"] for item in outside]
        blocks.insert(
            0,
            TextBlock(
                text="\n".join(lines),
                location=f"{location_prefix}, OCR-текст",
                block_type="ocr",
                metadata={"ocr": True, "confidence": round(float(np.mean(confidences)), 3) if confidences else 0.0},
            ),
        )
    return blocks


def _page_has_large_image(page) -> bool:
    page_area = max(1.0, float(page.rect.width * page.rect.height))
    try:
        for info in page.get_image_info():
            bbox = fitz.Rect(info["bbox"])
            if bbox.width * bbox.height / page_area >= 0.35:
                return True
    except (KeyError, RuntimeError, ValueError):
        return False
    return False


def _needs_ocr(page, text: str) -> bool:
    if settings.ocr_pdf_mode == "always":
        return True
    if settings.ocr_pdf_mode == "off":
        return False
    clean = "".join(character for character in text if character.isalnum())
    return len(clean) < 80 or _page_has_large_image(page)


def parse_pdf(path: Path) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text("text", sort=True).strip()
            if _needs_ocr(page, text):
                pix = page.get_pixmap(dpi=settings.ocr_dpi, alpha=False)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_blocks = _ocr_image_blocks(image, f"стр. {page_number}")
                if ocr_blocks:
                    blocks.extend(ocr_blocks)
                    continue
            if text:
                blocks.append(TextBlock(text=text, location=f"стр. {page_number}", metadata={"page": page_number}))
    return blocks


def parse_docx(path: Path) -> list[TextBlock]:
    doc = DocxDocument(path)
    blocks: list[TextBlock] = []
    section = "начало документа"
    paragraph_number = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph_number += 1
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "heading" in style or "заголовок" in style:
            section = text
        blocks.append(
            TextBlock(
                text=text,
                location=f"раздел «{section}», абзац {paragraph_number}",
                metadata={"section": section, "paragraph": paragraph_number},
            )
        )
    for table_number, table in enumerate(doc.tables, start=1):
        rows = []
        for row_number, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                rows.append((row_number, " | ".join(cells)))
        for start in range(0, len(rows), 15):
            group = rows[start : start + 15]
            if group:
                blocks.append(
                    TextBlock(
                        text="\n".join(text for _, text in group),
                        location=f"таблица {table_number}, строки {group[0][0]}–{group[-1][0]}",
                        block_type="table",
                        metadata={"table": table_number, "row_start": group[0][0], "row_end": group[-1][0]},
                    )
                )
    return blocks


def parse_doc(path: Path) -> list[TextBlock]:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        raise RuntimeError("Для чтения .doc требуется LibreOffice, но он не найден")
    with tempfile.TemporaryDirectory(prefix="doc-convert-") as folder:
        output_dir = Path(folder)
        result = subprocess.run(
            [executable, "--headless", "--convert-to", "docx", "--outdir", str(output_dir), str(path)],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
        )
        converted = output_dir / f"{path.stem}.docx"
        if result.returncode != 0 or not converted.exists():
            details = (result.stderr or result.stdout or "неизвестная ошибка").strip()
            raise RuntimeError(f"LibreOffice не смог преобразовать .doc: {details}")
        blocks = parse_docx(converted)
        html_result = subprocess.run(
            [executable, "--headless", "--convert-to", "html", "--outdir", str(output_dir), str(path)],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
        )
        html_path = output_dir / f"{path.stem}.html"
        if html_result.returncode == 0 and html_path.exists():
            tree = lxml_html.parse(str(html_path))
            table_blocks: list[TextBlock] = []
            compact_tables: list[str] = []
            existing_tables = {
                re.sub(r"\s+", "", block.text.replace("|", ""))
                for block in blocks
                if block.block_type == "table"
            }
            for table_number, table in enumerate(tree.xpath("//table"), start=1):
                rows = []
                for row_number, row in enumerate(table.xpath(".//tr"), start=1):
                    cells = [" ".join(cell.text_content().split()) for cell in row.xpath("./th|./td")]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    rendered = "\n".join(rows)
                    compact_rendered = re.sub(r"\s+", "", rendered.replace("|", ""))
                    compact_tables.append(compact_rendered)
                    if compact_rendered in existing_tables:
                        continue
                    table_blocks.append(
                        TextBlock(
                            text=rendered,
                            location=f"таблица {table_number}",
                            block_type="table",
                            metadata={"table": table_number, "legacy_doc": True},
                        )
                    )
            if table_blocks:
                filtered = []
                for block in blocks:
                    compact = re.sub(r"\s+", "", block.text)
                    if any(compact and compact in table_text for table_text in compact_tables):
                        continue
                    filtered.append(block)
                blocks = filtered + table_blocks
        return blocks


def _clean_cell(value) -> str:
    if pd.isna(value):
        return ""
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except (TypeError, ValueError):
            pass
    return " ".join(str(value).split())


def parse_xlsx(path: Path) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    book = pd.ExcelFile(path, engine="openpyxl")
    for sheet_name in book.sheet_names:
        frame = pd.read_excel(path, sheet_name=sheet_name, header=None, dtype=object, engine="openpyxl")
        frame = frame.dropna(how="all").dropna(axis=1, how="all")
        rows: list[tuple[int, str]] = []
        for index, values in frame.iterrows():
            cells = [_clean_cell(value) for value in values.tolist()]
            rendered = " | ".join(value for value in cells if value)
            if rendered:
                rows.append((int(index) + 1, rendered))
        for start in range(0, len(rows), 20):
            group = rows[start : start + 20]
            if not group:
                continue
            text = f"Лист: {sheet_name}\n" + "\n".join(f"Строка {n}: {value}" for n, value in group)
            blocks.append(
                TextBlock(
                    text=text,
                    location=f"лист «{sheet_name}», строки {group[0][0]}–{group[-1][0]}",
                    block_type="table",
                    metadata={"sheet": sheet_name, "row_start": group[0][0], "row_end": group[-1][0]},
                )
            )
    return blocks


def parse_image(path: Path) -> list[TextBlock]:
    with Image.open(path) as image:
        return _ocr_image_blocks(image, "изображение")


def parse_audio(path: Path) -> list[TextBlock]:
    segments, _ = _whisper_model().transcribe(str(path), beam_size=5, language="ru")
    text = " ".join(segment.text.strip() for segment in segments).strip()
    return [TextBlock(text=text, location="расшифровка аудио", block_type="audio")] if text else []


def parse_file(path: Path) -> list[TextBlock]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return parse_pdf(path)
    if extension == ".docx":
        return parse_docx(path)
    if extension == ".doc":
        return parse_doc(path)
    if extension == ".xlsx":
        return parse_xlsx(path)
    if extension in {".jpg", ".jpeg", ".png"}:
        return parse_image(path)
    if extension in {".mp3", ".wav", ".m4a", ".ogg", ".flac"}:
        return parse_audio(path)
    raise ValueError(f"Формат {extension or '<без расширения>'} не поддерживается")
